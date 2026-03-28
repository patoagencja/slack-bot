import asyncio
import os
from playwright.async_api import async_playwright
from docx import Document
from datetime import datetime

EMAIL = os.environ.get("DNA_EMAIL", "TWOJ_EMAIL")
PASSWORD = os.environ.get("DNA_PASSWORD", "TWOJE_HASLO")
LOGIN_URL = "https://dnarynkow.pl/wp-login.php"
CATEGORY_URL = "https://dnarynkow.pl/category/dna-premium/"

async def scrape():
    doc = Document()
    doc.add_heading("DNA Rynków – DNA Premium", 0)
    doc.add_paragraph(f"Pobrano: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=False)
        page = await browser.new_page()

        # Logowanie
        await page.goto(LOGIN_URL)
        await page.fill('#user_login', EMAIL)
        await page.fill('#user_pass', PASSWORD)
        await page.click('#wp-submit')
        await page.wait_for_load_state('networkidle')
        print("Zalogowano")

        # Zbieranie linków do artykułów
        article_links = []
        current_url = CATEGORY_URL
        while current_url:
            await page.goto(current_url)
            await page.wait_for_load_state('networkidle')
            links = await page.eval_on_selector_all(
                'article a[rel="bookmark"]',
                'els => [...new Set(els.map(e => e.href))]'
            )
            article_links.extend(links)
            print(f"Znaleziono {len(links)} artykułów na stronie: {current_url}")

            # Następna strona
            next_btn = await page.query_selector('a.next.page-numbers')
            current_url = await next_btn.get_attribute('href') if next_btn else None

        print(f"Łącznie artykułów: {len(article_links)}")

        # Scraping treści
        for i, url in enumerate(article_links):
            try:
                await page.goto(url)
                await page.wait_for_load_state('networkidle')
                title = await page.title()

                content_el = await page.query_selector('.entry-content, .post-content, article .content')
                content = await content_el.inner_text() if content_el else "Brak treści"

                date_el = await page.query_selector('time.entry-date, .post-date, time')
                date = await date_el.inner_text() if date_el else ""

                doc.add_heading(title, level=1)
                if date:
                    doc.add_paragraph(f"Data: {date}", style='Intense Quote')
                doc.add_paragraph(content)
                doc.add_paragraph("---")
                print(f"[{i+1}/{len(article_links)}] {title}")
            except Exception as e:
                print(f"Błąd przy {url}: {e}")

        await browser.close()

    filename = f"dna_premium_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    print(f"Zapisano: {filename}")

asyncio.run(scrape())
